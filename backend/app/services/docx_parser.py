"""
Extract text from DOCX files for the agents pipeline.
"""

import io
import logging
import docx

logger = logging.getLogger(__name__)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Parses a DOCX file from bytes and returns all extracted text from paragraphs.
    """
    doc = docx.Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs if p.text]
    full_text = "\n".join(parts).strip()
    logger.info("Extracted text from DOCX (%s paragraphs)", len(doc.paragraphs))
    return full_text
